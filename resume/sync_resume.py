#!/usr/bin/env python3
"""Bidirectional sync: cai_shengcheng_resume.docx <-> cai_shengcheng_resume.html"""

from __future__ import annotations

import argparse
import html
import re
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

RESUME_DIR = Path(__file__).resolve().parent
DOCX_PATH = RESUME_DIR / "cai_shengcheng_resume.docx"
HTML_PATH = RESUME_DIR / "cai_shengcheng_resume.html"

NAVY = (0, 51, 102)
GRAY = (102, 102, 102)
DARK_GRAY = (85, 85, 85)

SECTION_HEADINGS = {"教育背景", "工作经历", "专业技能"}


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def parse_html_runs(fragment: str) -> list[tuple[str, bool]]:
    runs: list[tuple[str, bool]] = []

    class Parser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.stack: list[bool] = [False]

        def handle_starttag(self, tag, attrs):
            if tag == "strong":
                self.stack.append(True)

        def handle_endtag(self, tag):
            if tag == "strong" and len(self.stack) > 1:
                self.stack.pop()

        def handle_data(self, data):
            text = data.replace("\xa0", " ")
            if text:
                runs.append((text, self.stack[-1]))

    parser = Parser()
    parser.feed(fragment)
    return runs


def add_runs_to_paragraph(paragraph, fragment: str, size=10.5, bold_color=NAVY):
    for text, bold in parse_html_runs(fragment):
        run = paragraph.add_run(text)
        set_run_font(run, size, bold, bold_color if bold else None)


def para_to_html_runs(paragraph) -> str:
    parts: list[str] = []
    for run in paragraph.runs:
        text = escape_html(run.text)
        if not text:
            continue
        if run.bold:
            parts.append(f"<strong>{text}</strong>")
        else:
            parts.append(text)
    return "".join(parts).strip()


def job_text_to_html(text: str) -> str:
    parts = re.split(r"\s{2,}", text.strip())
    if len(parts) >= 3:
        company, date, role = parts[0], parts[1], parts[2]
        return (
            f"<strong>{escape_html(company)}</strong>"
            f"&nbsp;&nbsp;&nbsp;&nbsp;"
            f'<span class="date">{escape_html(date)}</span>'
            f"&nbsp;&nbsp;&nbsp;&nbsp;"
            f'<span class="role">{escape_html(role)}</span>'
        )
    return escape_html(text)


Block = tuple[str, str]


# ---------------------------------------------------------------------------
# HTML -> blocks
# ---------------------------------------------------------------------------

def extract_html_blocks(html_text: str) -> list[Block]:
    body_match = re.search(r"<body>(.*)</body>", html_text, re.S)
    if not body_match:
        raise ValueError("No <body> found in HTML")
    body = body_match.group(1)

    patterns = [
        (r"<h1[^>]*>(.*?)</h1>", "h1"),
        (r'<div class="meta">(.*?)</div>', "meta"),
        (r'<div class="contact">(.*?)</div>', "contact"),
        (r'<div class="summary">\s*(.*?)\s*</div>', "summary"),
        (r'<h2 class="section-heading">(.*?)</h2>', "h2"),
        (r'<div class="edu-item">(.*?)</div>', "edu"),
        (r'<div class="job-line">\s*(.*?)\s*</div>', "job"),
        (r"<h3>(.*?)</h3>", "h3"),
        (r"<li>(.*?)</li>", "li"),
        (r'<p class="minor">(.*?)</p>', "minor"),
        (r'<div class="skill-line">(.*?)</div>', "skill"),
    ]

    blocks: list[Block] = []
    pos = 0
    while pos < len(body):
        best = None
        for pattern, kind in patterns:
            match = re.search(pattern, body[pos:], re.S)
            if match and (best is None or match.start() < best[0]):
                best = (match.start(), match.end(), kind, match.group(1))
        if not best:
            break
        start, end, kind, inner = best
        pos += start
        inner = html.unescape(inner.strip())
        if "header-photo" in inner or "<img" in inner:
            pos += end - start
            continue
        blocks.append((kind, inner))
        pos += end - start
    return blocks


# ---------------------------------------------------------------------------
# DOCX -> blocks
# ---------------------------------------------------------------------------

def extract_docx_items(doc: Document) -> list[dict]:
    items: list[dict] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        items.append(
            {
                "style": paragraph.style.name if paragraph.style else "Normal",
                "text": text,
                "html": para_to_html_runs(paragraph),
            }
        )
    return items


def extract_docx_blocks(doc: Document) -> list[Block]:
    items = extract_docx_items(doc)
    if len(items) < 8:
        raise ValueError("Unexpected docx structure: too few paragraphs")

    blocks: list[Block] = []
    idx = 0

    def take(kind: str, *, use_text: bool = False):
        nonlocal idx
        content = items[idx]["text"] if use_text else items[idx]["html"]
        blocks.append((kind, content))
        idx += 1

    take("h1")
    take("meta")
    take("contact")
    take("summary")
    take("h2", use_text=True)
    take("edu")
    take("edu")
    take("h2", use_text=True)
    blocks.append(("job", job_text_to_html(items[idx]["text"])))
    idx += 1

    while idx < len(items):
        item = items[idx]
        text = item["text"]
        style = item["style"]

        if text in SECTION_HEADINGS:
            blocks.append(("h2", text))
            idx += 1
            continue
        if style == "List Bullet":
            blocks.append(("li", item["html"]))
            idx += 1
            continue
        if text.startswith("早期曾参与"):
            blocks.append(("minor", item["html"]))
            idx += 1
            continue
        if text.startswith("编程基础：") or text.startswith("推理优化：") or text.startswith("算子开发：") or text.startswith("开发技能："):
            blocks.append(("skill", item["html"]))
            idx += 1
            continue
        blocks.append(("h3", item["html"]))
        idx += 1

    return blocks


# ---------------------------------------------------------------------------
# blocks -> DOCX
# ---------------------------------------------------------------------------

def build_docx(blocks: list[Block]) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    for kind, inner in blocks:
        if kind == "h1":
            paragraph = doc.add_paragraph()
            add_runs_to_paragraph(paragraph, inner, 20, NAVY)
        elif kind == "meta":
            paragraph = doc.add_paragraph()
            add_runs_to_paragraph(paragraph, inner, 9.5, GRAY)
            paragraph.paragraph_format.space_after = Pt(2)
        elif kind == "contact":
            paragraph = doc.add_paragraph()
            add_runs_to_paragraph(paragraph, inner, 9.5, GRAY)
            paragraph.paragraph_format.space_after = Pt(2)
        elif kind == "summary":
            paragraph = doc.add_paragraph()
            add_runs_to_paragraph(paragraph, inner, 10)
            paragraph.paragraph_format.space_after = Pt(8)
        elif kind == "h2":
            paragraph = doc.add_paragraph()
            add_runs_to_paragraph(paragraph, inner, 11, NAVY)
            paragraph.paragraph_format.space_after = Pt(4)
        elif kind == "edu":
            paragraph = doc.add_paragraph()
            add_runs_to_paragraph(paragraph, inner, 9.5)
        elif kind == "job":
            paragraph = doc.add_paragraph()
            add_runs_to_paragraph(paragraph, inner, 11)
        elif kind == "h3":
            paragraph = doc.add_paragraph()
            add_runs_to_paragraph(paragraph, inner, 10.5, NAVY)
            paragraph.paragraph_format.space_before = Pt(6)
        elif kind == "li":
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
            add_runs_to_paragraph(paragraph, inner, 10.5)
        elif kind == "minor":
            paragraph = doc.add_paragraph()
            add_runs_to_paragraph(paragraph, inner, 9.5, DARK_GRAY)
            paragraph.paragraph_format.space_before = Pt(6)
        elif kind == "skill":
            paragraph = doc.add_paragraph()
            add_runs_to_paragraph(paragraph, inner, 10)
            paragraph.paragraph_format.space_after = Pt(2)

    return doc


# ---------------------------------------------------------------------------
# blocks -> HTML body
# ---------------------------------------------------------------------------

def strip_html_tags(fragment: str) -> str:
    return re.sub(r"<[^>]+>", "", fragment)


def render_html_body(blocks: list[Block]) -> str:
    lines: list[str] = []
    lines.extend(
        [
            "    <div class=\"header-top\">",
            "      <div class=\"header-info\">",
        ]
    )

    idx = 0
    while idx < len(blocks) and blocks[idx][0] in {"h1", "meta", "contact"}:
        kind, content = blocks[idx]
        if kind == "h1":
            lines.append(f"        <h1>{strip_html_tags(content)}</h1>")
        elif kind == "meta":
            lines.append(f'        <div class="meta">{content}</div>')
        elif kind == "contact":
            lines.append(f'        <div class="contact">{content}</div>')
        idx += 1

    lines.extend(
        [
            "      </div>",
            "      <div class=\"header-photo\">",
            "        <img src=\"csc.jpg\" alt=\"蔡圣诚证件照\">",
            "      </div>",
            "    </div>",
            "",
        ]
    )

    while idx < len(blocks):
        kind, content = blocks[idx]
        if kind == "summary":
            lines.append(f'    <div class="summary">\n      {content}\n    </div>\n')
        elif kind == "h2":
            lines.append(f'    <h2 class="section-heading">{strip_html_tags(content)}</h2>')
        elif kind == "edu":
            lines.append(f'    <div class="edu-item">{content}</div>')
        elif kind == "job":
            lines.append(f'    <div class="job-line">\n      {content}\n    </div>\n')
        elif kind == "h3":
            lines.append(f"    <h3>{content}</h3>")
        elif kind == "li":
            ul_lines = ["    <ul>"]
            while idx < len(blocks) and blocks[idx][0] == "li":
                ul_lines.append(f"      <li>{blocks[idx][1]}</li>")
                idx += 1
            ul_lines.append("    </ul>")
            lines.append("\n".join(ul_lines))
            continue
        elif kind == "minor":
            lines.append(f'    <p class="minor">{content}</p>')
        elif kind == "skill":
            lines.append(f'    <div class="skill-line">{content}</div>')
        idx += 1

    return "\n".join(lines) + "\n"


def rebuild_html(html_text: str, blocks: list[Block]) -> str:
    body_inner = render_html_body(blocks)
    return re.sub(
        r"(<body>\s*<div class=\"page\">).*?(</div>\s*</body>)",
        rf"\1\n{body_inner}  \2",
        html_text,
        count=1,
        flags=re.S,
    )


# ---------------------------------------------------------------------------
# Sync entry points
# ---------------------------------------------------------------------------

def sync_html_to_docx():
    html_text = HTML_PATH.read_text(encoding="utf-8")
    blocks = extract_html_blocks(html_text)
    doc = build_docx(blocks)
    doc.save(DOCX_PATH)
    print(f"html2docx: synced {len(blocks)} blocks -> {DOCX_PATH}")


def sync_docx_to_html():
    html_text = HTML_PATH.read_text(encoding="utf-8")
    doc = Document(DOCX_PATH)
    blocks = extract_docx_blocks(doc)
    updated = rebuild_html(html_text, blocks)
    HTML_PATH.write_text(updated, encoding="utf-8")
    print(f"docx2html: synced {len(blocks)} blocks -> {HTML_PATH}")


def main():
    parser = argparse.ArgumentParser(description="Sync resume between DOCX and HTML")
    parser.add_argument(
        "direction",
        choices=["html2docx", "docx2html"],
        help="html2docx: HTML -> DOCX; docx2html: DOCX -> HTML",
    )
    args = parser.parse_args()
    if args.direction == "html2docx":
        sync_html_to_docx()
    else:
        sync_docx_to_html()


if __name__ == "__main__":
    main()
