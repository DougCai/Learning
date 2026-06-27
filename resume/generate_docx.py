#!/usr/bin/env python3
"""Generate resume DOCX from structured content."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

OUT = "/home/caishengcheng/resume/cai_shengcheng_resume.docx"

def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, 18, True, (0, 51, 102))
    else:
        set_run_font(run, 11, True, (0, 51, 102))
        p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, bold_part, normal_part):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(bold_part)
    set_run_font(r1, 10.5, True, (0, 51, 102))
    r2 = p.add_run(normal_part)
    set_run_font(r2, 10.5)

def main():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    p = doc.add_paragraph()
    r = p.add_run("蔡圣诚")
    set_run_font(r, 20, True, (0, 51, 102))

    for line in [
        "求职方向：大模型推理工程师（量化推理 / 推理性能优化）",
        "杭州（可接受上海） | 15958033960 | cscznb5612@gmail.com",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_run_font(r, 9.5, color=(102, 102, 102))
        p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run(
        "4 年华为昇腾大模型推理研发经验，现任下一代模型量化工具负责人。主导 MXFP8/MXFP4 与 W4A4 "
        "低精度量化算法及推理落地，支撑 DeepSeek v4、Qwen3.5 等 20+ 权重在昇腾 A5 上线与客户送测；"
        "精度误差控制在 1% 以内，推理性能达竞品 1.x 倍，单模型适配周期压缩至 1 天。"
        "协同 4 部门 20+ 人完成推理引擎、量化算法与工具链端到端交付。"
    )
    set_run_font(r, 10)
    p.paragraph_format.space_after = Pt(8)

    add_heading(doc, "专业技能", 2)
    skills = [
        ("编程基础", "Python、C/C++（11 标准及以上），熟悉常用数据结构与面向对象设计"),
        ("深度学习", "PyTorch 大模型推理 Pipeline 构建，核心模块 PTQ 量化与性能调优"),
        ("量化算法", "MXFP8/MXFP4、SVDQuant、二级量化、MSE Round、W4A4 PTQ；熟悉 AWQ/GPTQ/SmoothQuant 等 PTQ 范式"),
        ("推理优化", "量化算子实现、Model-free 权重转换、推理引擎协同优化、0-day 模型适配"),
        ("通信与算子", "多卡通信算子调试（Allreduce 等）、Ascend C 算子开发与上板调试"),
    ]
    for label, val in skills:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}：")
        set_run_font(r1, 10, True, (0, 51, 102))
        r2 = p.add_run(val)
        set_run_font(r2, 10)
        p.paragraph_format.space_after = Pt(2)

    add_heading(doc, "工作经历", 2)
    p = doc.add_paragraph()
    r1 = p.add_run("华为技术有限公司")
    set_run_font(r1, 11, True)
    r2 = p.add_run("    2022.02 — 至今")
    set_run_font(r2, 10, color=(102, 102, 102))
    p = doc.add_paragraph()
    r = p.add_run("高级工程师 · 昇腾大模型推理 / 量化")
    set_run_font(r, 10, color=(85, 85, 85))

    for title, bullets in [
        ("大模型量化推理（项目负责人）", [
            ("MXFP 量化落地：", "主导 MXFP8、MXFP4 算法设计与实现，支撑 Qwen、DeepSeek、GLM、LongCat、Kimi、Minimax 等 20+ 权重在昇腾 A5 NPU 上线及客户送测；对 DeepSeek v4、Qwen3.5 实现 0-day 量化适配。精度相对浮点误差 <1%，推理性能达竞品 1.x 倍，单模型适配 + 精度调优周期约 1 天。"),
            ("W4A4 低精度量化：", "主导 SVDQuant、二级量化、MSE Round、C7 等 4bit 方案，在 Wan2.2 上完成 W4A4 量化；VBench 精度指标误差均 <0.01，推理性能达竞品 1.x 倍，相对浮点精度损失 <1%。"),
            ("动态 PTQ 转换：", "实现 Model-free 权重转换能力，在 PTQ 动态场景下以轻量化流程完成权重量化，降低大模型迁移与推理部署成本。"),
            ("端到端项目交付：", "作为负责人协同推理引擎、量化算法、量化工具 4 个部门 20+ 人推进交付，保障客户面送测通过；相较上一代引入浮点量化体系并扩展 4bit 低精度算法矩阵，参与关键方案设计与路线决策。"),
        ]),
        ("昇腾算子调试工具", [
            ("上板调试与生态：", "基于 LLDB 实现 NPU 断点、单步、变量打印及 coredump / 多 Kernel / 多卡调试，端到端打通真机调试流程；社区覆盖 1000+ 开发人员、20+ 活跃算子团队，算子精度与挂死类问题定位由数天缩短至小时级，开发效率提升数倍。"),
        ]),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(title)
        set_run_font(r, 10.5, True)
        p.paragraph_format.space_before = Pt(6)
        for b, n in bullets:
            add_bullet(doc, b, n)

    p = doc.add_paragraph()
    r = p.add_run("早期参与农机自动驾驶路径规划算法研发（专利 1 项）。")
    set_run_font(r, 9.5, color=(85, 85, 85))
    p.paragraph_format.space_before = Pt(6)

    add_heading(doc, "教育背景", 2)
    for line in [
        "华东师范大学 · 计算机技术 · 硕士 · 2019 — 2022",
        "浙江工业大学 · 计算机科学与技术 · 本科 · 2015 — 2019",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_run_font(r, 10.5)

    add_heading(doc, "专利", 2)
    p = doc.add_paragraph()
    r = p.add_run("发明专利 1 项（耕地路径规划）")
    set_run_font(r, 10.5)

    doc.save(OUT)
    print(f"Saved: {OUT}")

if __name__ == "__main__":
    main()
